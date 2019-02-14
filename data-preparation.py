#!/usr/bin/env python
# encoding: utf-8

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
from docx import Document
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
import re
import pandas as pd
import sys


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                table_res = paragraph.text
        print("\n")


### update the headers to the existing dictionary
def iter_heading(paragraphs):
    for paragraph in paragraphs:
        item = paragraph.style.name
        item_list = ['Heading 1', 'Heading 2', 'Heading 3']
        if item in item_list:
            yield (paragraph, item)


def get_headers(file_path):
    document = Document(file_path)
    headers_list = []
    for heading in iter_heading(document.paragraphs):
        headers = (heading[0].text).strip()
        headers_list.append((headers, heading[1]))
    return headers_list


def tag_documents(file_path):

    document = Document(file_path)
    cnt = 0
    results_para = {}
    for block in iter_block_items(document):
        cnt += 1
        if isinstance(block, Paragraph):
            raw_text = block.text
            raw_text = raw_text.strip()

            if len(raw_text) > 1:

                if len(raw_text) < 150 and len(raw_text) > 30:
                    raw_text = [' <li> ', raw_text]
                    results_para.update({cnt: raw_text})

                elif len(raw_text) < 30:
                    raw_text = ['<heading> ', raw_text]
                    results_para.update({cnt: raw_text})

                elif len(raw_text) > 150:
                    raw_text = [' <para> ', raw_text]
                    results_para.update({cnt: raw_text})

                elif isinstance(block, Table):
                    res = table_print(block)
                    res = [' <TABLE> ', res]
                    results_para.update({cnt: res})
            else:
                res = [' < None > ', 'None']
                results_para.update({cnt: res})
        else:
            res = [' < None > ', 'None']
            results_para.update({cnt: res})
    return results_para


def update_headers(results_para, headers_list):
    dict_res = {}
    for k, v in results_para.items():
        for h in headers_list:
            if h[0] == v[1]:
                new_value = [(h[1], v[0]), h[0]]
                dict_res[k] = new_value
            else:
                pass
    results_para.update(dict_res)
    return results_para


def label_documents(df):
    df.assign(new_heading="")
    t_text = ""
    for i, row in df.iterrows():
        if len(row['heading_type']) == 2:
            t_text = ""
            t_text = row['para_text']
            row['new_heading'] = t_text
        row['new_heading'] = t_text
    return df


if __name__ == "__main__":

    file_path = str(sys.argv[1])
    results_para = tag_documents(file_path)
    headers_list = get_headers(file_path)
    results_para = update_headers(results_para, headers_list)

    data_df = pd.DataFrame.from_dict(results_para, orient='index', columns=['heading_type', 'para_text'])
    final_result = data_df[data_df['para_text'] != 'None']
    final_res = label_documents(final_result)

    final_res.to_csv('./out_file/OUT.csv')
