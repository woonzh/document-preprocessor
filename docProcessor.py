#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Jan 22 16:52:27 2019

@author: zhenhao
"""

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pandas as pd
import argparse

filePath='./data/nems_111010.docx'
filePath='./data/sample.docx'
filePath='./data/Telekom malaysia chatbot Proposal Draft.docx'
#filePath='./data/017921w14 Pattern Manager 6.0 - Operating User Guide.docx'
#filePath='./data/Unstructured_Assurance of DOT and TBiz Voice on nbn FTTP.docx'    

class utils:
    def __init__(self):
        self.ranking={
            'Title': 1,
            'Subtitle': 2,
            'Heading': 3,
            'Body Text':4,
            'Normal': 4,
            'List': 4,
            'Intense Quote': 4,
            'Caption':4,
            'Footer': 4,
            'Quote': 4,
            'TOC':0
            }
    
    def assignStyleRank(self, name):
        for i in self.ranking:
            if i in name:
                nameSplit=name.split()
                for j in nameSplit:
                    try:
                        index=int(j)
                    except:
                        index=0
                styleRank=self.ranking[i]*10 + index
                return styleRank, self.ranking[i]
        
        return 40, 4

class docs:
    def __init__(self,text, rank, bold, name):
        self.text=text
        self.rank=rank
        if bold=='True':
            self.bold=1
        else:
            self.bold=0
        self.length=len(self.text)
        self.child=[]
        self.parent=None
        self.name=name
        self.styleRank, self.styleRankGen=utils().assignStyleRank(name)
        
    def assignChild(self,ele):
        self.child.append(ele)
    
    def assignParent(self,ele):
        self.parent=ele
        
class docStruct:
    def __init__(self):
        self.parent=docs('parent',0, 1, 'Title')
        self.store=[self.parent]
        self.prev=self.parent
        self.cur=self.parent
        self.dictStore=None
        
    def allocate(self, curCompare):
        if curCompare == 'smaller':
            self.cur.assignParent(self.prev)
            self.prev.assignChild(self.cur)
            self.prev=self.cur
        
        if curCompare == 'same':
            self.cur.assignParent(self.prev.parent)
            self.prev.parent.assignChild(self.cur)
            self.prev=self.cur
            
        if curCompare == 'bigger':
            self.prev=self.prev.parent
            self.findParents() 
        
        
    def findParents(self):
        #iterate through nodes and find relationship
        #if current is smaller, assign prev as parent and cur as child
        # if same, assign the parent of the prev as the parent of the cur and assign cur as child of the parent of prev
        # if current is bigger, repeat the above but with prev as the parent of the prev
        if self.prev.styleRankGen < self.cur.styleRankGen:
            self.allocate('smaller')
        else:
            if self.prev.styleRankGen == self.cur.styleRankGen:
#                self.allocate('same')
                if self.prev.rank<self.cur.rank or self.prev.bold>self.cur.bold:
                    self.allocate('smaller')
                else:
                    if self.prev.rank==self.cur.rank:
                        self.allocate('same')
                    else:
                        self.allocate('bigger')
            else:
                self.allocate('bigger')
    
    def addNewNode(self, text, rank, bold, name):
        self.cur=docs(text, rank, bold, name)
        self.store.append(self.cur)
        if self.cur.styleRank >= 10:
            self.findParents()
        
    def compileLoop(self,head):
        if len(head.child)==0:
            return head.text
        store=[]
        for i in head.child:
            store.append(self.compileLoop(i))
        
        return {head.text:store}
        
    def compileDict(self):
        store=self.compileLoop(self.parent)
        return store
    
    def compileHeaderTextLoop(self, start):
        name=str(list(start)[0])
        start=start[name]
        store=[]
        text=''
        for i in start:
            if isinstance(i, str):
                store.append(i)
                text=text+'\n' +i
            if isinstance(i, dict):
                self.compileHeaderTextLoop(i)
        self.headerTextCompilation.append({name:[store, text]})
        self.headerTextDF.loc[len(self.headerTextDF)]=[name, text]
    
    def compileHeaderText(self):
        self.headerTextCompilation=[]
        self.headerTextDF=pd.DataFrame(columns=['header', 'text'])
        self.compileHeaderTextLoop(self.dictStore)
    
    def processDF(self,df):
        for i in range(len(df)):
            self.addNewNode(df.iloc[i,0], df.iloc[i,7], df.iloc[i,1], df.iloc[i,5])
        
        self.dictStore=self.compileDict()
        self.compileHeaderText()
        
def variableGetter(curList, nextList):
    finalList=[]
    for count, rec in enumerate(curList):
        if rec!='None':
            finalList.append(rec)
        else:
            finalList.append(nextList[count])
    
    return finalList

def fontExtractor(paragraph, doc):
    ans=[]
    ans2=[]
    
    runs=paragraph.runs
    
    headName=paragraph.style.name
    
    if str(paragraph.style.font.size) == 'None':
        paragraph.style=list(doc.styles)[0]
    
    for i in range(len(runs)):
        curRun=runs[i].font
        ans.append(str(paragraph.text))
        ans.append(str(curRun.bold))
        ans.append(str(curRun.italic))
        ans.append(str(curRun.underline))
        try:
            ans.append(curRun.size.pt)
        except:
            ans.append(1000)
        ans.append('None')
        ans.append(str(paragraph.alignment))
        if i==0:
            finalAns=variableGetter(ans,ans)
        else:
            finalAns=variableGetter(finalAns, ans)
    
    style=paragraph.style.font
    ans2.append(str(paragraph.text))
    ans2.append(str(style.bold))
    ans2.append(str(style.italic))
    ans2.append(str(style.underline))
    try:
        ans2.append(style.size.pt)
    except:
        ans2.append(1000)
    ans2.append(headName)
    ans2.append(str(paragraph.alignment))
    
    finalAns=variableGetter(finalAns, ans2)
    
    styleRank, styleRankGen = utils().assignStyleRank(finalAns[5])
    finalAns.append(styleRank)
    finalAns.append(styleRankGen)
    
    return finalAns

def docProfileCreator(doc):
    df=pd.DataFrame(columns=['text', 'bold', 'italic', 'underline', 'size', 'name', 'alignment', 'style rank', 'style rank general'])
    for count, i in enumerate(list(doc.paragraphs)):
        if i.text!='':
            df.loc[count]=fontExtractor(i, doc)
            
    sizeList=list(df['size'])
    minSize=min(sizeList)
    newSizeList=[minSize if x==1000 else x for x in sizeList]
    df['size']=newSizeList
            
    sizes=[int(i) for i in set(df['size'])]
    sizes.sort(reverse=True)
    sizeVal={val:count+1 for count, val in enumerate(sizes)}
    
    df['rank']=[sizeVal[int(val)] for val in df['size']]
        
    return df

def findParents(ele, prev):
    if prev.rank>ele.rank:
        ele.assignParent(prev)
        prev.assignChild(ele)
        return prev
    if prev.rank==ele.rank:
        ele.assignParent(prev.parent)
        prev.parent.assignChild(ele)
        return ele
    if prev.rank<ele.rank:
        findParents(ele, prev.parent)

#call this to get the relationship tree
def getDocTree(filePath):
    doc=Document(filePath)
    df=docProfileCreator(doc)
    
    docStore=docStruct()
    docStore.processDF(df)
    docDict=docStore.dictStore
    return docDict

# call this to generate textCSV
def generateHeaderTextCSV(filePath):
    doc=Document(filePath)
    df=docProfileCreator(doc)
    
    docStore=docStruct()
    docStore.processDF(df)
    docDict=docStore.dictStore
    docTable=docStore.headerTextCompilation
    docDF=docStore.headerTextDF
    docDF.to_csv('header_text.csv')

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Process docx files')
    parser.add_argument("-f", dest='filePath', action="store", default='./data/Telekom malaysia chatbot Proposal Draft.docx', help='input file path')
    args = parser.parse_args()
    
    generateHeaderTextCSV(args.filePath)
